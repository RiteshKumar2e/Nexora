"""Parsers for uploaded documents (PDF, DOCX, CSV, Excel, TXT, JSON, code)."""
from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path


def parse_txt(file_bytes: bytes) -> str:
    """Parse raw text files."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="replace")


def parse_csv(file_bytes: bytes) -> str:
    """Parse CSV and convert to Markdown table format for LLM readability."""
    text = parse_txt(file_bytes)
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if not lines:
        return ""
    
    markdown_table = []
    for i, line in enumerate(lines):
        row = [cell.strip() for cell in line.split(",")]
        # Escape markdown pipe
        escaped_row = [c.replace("|", "\\|") for c in row]
        markdown_table.append("| " + " | ".join(escaped_row) + " |")
        if i == 0:
            # Add separator line
            markdown_table.append("|" + "---|"*len(row))
            
    return "\n".join(markdown_table)


def parse_json(file_bytes: bytes) -> str:
    """Parse JSON and return clean formatted JSON string."""
    try:
        data = json.loads(parse_txt(file_bytes))
        return json.dumps(data, indent=2)
    except Exception as e:
        return f"Invalid JSON file: {e}"


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using pypdf (real extraction, page by page)."""
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(file_bytes))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        if txt.strip():
            pages.append(f"--- Page {i + 1} ---\n{txt.strip()}")
    if not pages:
        return "[PDF contained no extractable text — it may be a scanned image.]"
    return "\n\n".join(pages)


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a Word .docx using python-docx (paragraphs + tables)."""
    import docx

    doc = docx.Document(BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts) or "[DOCX contained no readable text.]"


def parse_xlsx(file_bytes: bytes) -> str:
    """Extract cell values from an Excel workbook (best effort; needs openpyxl)."""
    try:
        import openpyxl
    except ImportError:
        return "[XLSX parsing needs openpyxl: pip install openpyxl]"
    wb = openpyxl.load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"--- Sheet: {ws.title} ---")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_document(filename: str, file_bytes: bytes) -> str:
    """Dispatch file bytes to the right text extractor based on file extension."""
    ext = Path(filename).suffix.lower()

    if ext in (".txt", ".md", ".markdown", ".py", ".js", ".ts", ".tsx", ".jsx",
               ".html", ".htm", ".css", ".rs", ".go", ".c", ".cpp", ".java",
               ".rb", ".php", ".sh", ".sql", ".yaml", ".yml", ".xml", ".log",
               ".ini", ".toml"):
        return parse_txt(file_bytes)
    if ext == ".csv":
        return parse_csv(file_bytes)
    if ext == ".json":
        return parse_json(file_bytes)
    if ext == ".pdf":
        try:
            return parse_pdf(file_bytes)
        except Exception as exc:  # noqa: BLE001
            return f"[Could not read PDF: {exc}]"
    if ext in (".docx", ".doc"):
        try:
            return parse_docx(file_bytes)
        except Exception as exc:  # noqa: BLE001
            return f"[Could not read Word document: {exc}]"
    if ext in (".xlsx", ".xlsm"):
        try:
            return parse_xlsx(file_bytes)
        except Exception as exc:  # noqa: BLE001
            return f"[Could not read spreadsheet: {exc}]"
    # Unknown/binary: try a best-effort text decode.
    return parse_txt(file_bytes)[:50000]
